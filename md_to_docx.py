"""
Convert research_paper.md to a formatted DOCX file.
Uses python-docx to create a professional-looking Word document
with proper headings, tables, images, and formatting.
"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Paths
PROJECT = Path("/Users/abhishekdeep/Documents/Data/MTech/Project-Sem 2/RResearch_Project")
MD_FILE = Path("/Users/abhishekdeep/.gemini/antigravity/brain/8536dc05-73d5-4c06-bf91-5e10609f69bd/research_paper.md")
OUT_FILE = PROJECT / "research_paper.docx"
FIGURES_DIR = Path("/Users/abhishekdeep/.gemini/antigravity/brain/8536dc05-73d5-4c06-bf91-5e10609f69bd/figures")
PROJECT_FIGURES = PROJECT / "reports" / "figures"

def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = cell._tc.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)

def add_formatted_run(paragraph, text):
    """Add text with inline bold/italic/code formatting."""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            inner = part[2:-2]
            # Check for nested code
            code_parts = re.split(r'(`[^`]+`)', inner)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.bold = True
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    run = paragraph.add_run(cp)
                    run.bold = True
        elif '`' in part:
            # Inline code
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    run = paragraph.add_run(cp)
        else:
            paragraph.add_run(part)

def parse_table(lines):
    """Parse markdown table lines into list of rows (each row is list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:|]+\|$', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows

def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Clean markdown formatting
                clean = cell_text.replace('**', '').replace('*', '').replace('`', '')
                clean = re.sub(r'\\textbf\{([^}]*)\}', r'\1', clean)
                if i == 0:
                    # Header row
                    run = p.add_run(clean)
                    run.bold = True
                    run.font.size = Pt(9)
                    set_cell_shading(cell, "2F5496")
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run = p.add_run(clean)
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing

def add_image(doc, img_path, caption="", width=5.5):
    """Add an image with caption."""
    resolved = None
    if os.path.exists(img_path):
        resolved = img_path
    else:
        # Try figures directory
        basename = os.path.basename(img_path)
        for d in [FIGURES_DIR, PROJECT_FIGURES]:
            candidate = d / basename
            if candidate.exists():
                resolved = str(candidate)
                break
    
    if resolved and os.path.exists(resolved):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(resolved, width=Inches(width))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.italic = True
    else:
        p = doc.add_paragraph(f"[Figure: {caption or img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

def convert_md_to_docx():
    print("Reading research_paper.md...")
    with open(MD_FILE, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    doc = Document()
    
    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # -- Heading styles --
    for level in range(1, 4):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Times New Roman'
        hstyle.font.color.rgb = RGBColor(47, 84, 150)
        if level == 1:
            hstyle.font.size = Pt(16)
        elif level == 2:
            hstyle.font.size = Pt(13)
        else:
            hstyle.font.size = Pt(11)
    
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip horizontal rules
        if stripped == '---' or stripped == '***':
            i += 1
            continue
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # --- TITLE (# heading) ---
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            # Add as title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(47, 84, 150)
            
            # Add author
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run("Abhishek Deep")
            run2.font.size = Pt(12)
            run2.italic = True
            
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run("Department of Computer Science and Engineering\nIndian Institute of Technology")
            run3.font.size = Pt(10)
            
            i += 1
            continue
        
        # --- HEADINGS ---
        if stripped.startswith('#### '):
            heading_text = stripped[5:].strip()
            heading_text = re.sub(r'\*\*([^*]*)\*\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=4)
            i += 1
            continue
        if stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            heading_text = re.sub(r'\*\*([^*]*)\*\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            heading_text = re.sub(r'\d+\.\s*', '', heading_text, count=1)
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # --- IMAGES ---
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            caption = img_match.group(1)
            path = img_match.group(2)
            add_image(doc, path, caption)
            i += 1
            continue
        
        # --- TABLES ---
        if stripped.startswith('|'):
            # Collect table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table(doc, rows)
            continue
        
        # --- BLOCKQUOTES ---
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            quote_text = re.sub(r'\*\*([^*]*)\*\*', r'\1', quote_text)
            quote_text = re.sub(r'\*\[!.*?\]\*\s*', '', quote_text)
            quote_text = re.sub(r'\[!.*?\]\s*', '', quote_text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue
        
        # --- NUMBERED LISTS ---
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
            text = text.replace('`', '')
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text)
            i += 1
            continue
        
        # --- BULLET LISTS ---
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
            text = text.replace('`', '')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            i += 1
            continue
        
        # --- MATH BLOCKS ---
        if stripped.startswith('$$'):
            math_text = ''
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                math_text += lines[i].strip() + ' '
                i += 1
            i += 1  # skip closing $$
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text.strip())
            run.font.name = 'Cambria Math'
            run.italic = True
            run.font.size = Pt(11)
            continue
        
        # --- REGULAR PARAGRAPH ---
        # Clean markdown
        text = stripped
        # Remove inline math $ markers but keep content
        text = re.sub(r'\$([^$]+)\$', r'\1', text)
        
        p = doc.add_paragraph()
        add_formatted_run(p, text)
        i += 1
    
    # Save
    doc.save(str(OUT_FILE))
    print(f"\nSaved: {OUT_FILE}")
    print(f"Size: {os.path.getsize(OUT_FILE) / 1024:.0f} KB")

if __name__ == '__main__':
    convert_md_to_docx()
